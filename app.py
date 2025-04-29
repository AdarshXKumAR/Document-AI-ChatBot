
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename
import os
from pathlib import Path
import google.generativeai as genai
from groq import Groq
from llama_index.core import (
    VectorStoreIndex,
    Document,
)
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
import pypdf
import docx
import pandas as pd
import tempfile
from dotenv import load_dotenv

# Load environment variables from .env
load_dotenv(".env.local")

# Retrieve API keys from .env
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")    
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Ensure API keys are set
if not GEMINI_API_KEY or not GROQ_API_KEY:
    raise ValueError("API keys are missing. Please check your .env file.")

# Initialize APIs
genai.configure(api_key=GEMINI_API_KEY)
groq_client = Groq(api_key=GROQ_API_KEY)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Initialize embedding model
embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-mpnet-base-v2")

# Ensure upload directory exists
Path(app.config['UPLOAD_FOLDER']).mkdir(parents=True, exist_ok=True)

def process_file(file_path):
    """Process different file types and return text content"""
    file_extension = Path(file_path).suffix.lower()
    
    if file_extension == '.pdf':
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            text = ' '.join(page.extract_text() for page in pdf_reader.pages)
    
    elif file_extension == '.docx':
        doc = docx.Document(file_path)
        text = ' '.join(paragraph.text for paragraph in doc.paragraphs)
    
    elif file_extension in ['.xlsx', '.xls']:
        df = pd.read_excel(file_path)
        text = df.to_string()
    
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
    
    return text

def query_gemini(text, question):
    """Query using Google's Gemini API"""
    model = genai.GenerativeModel('gemini-pro')
    prompt = f"Context: {text}\n\nQuestion: {question}\n\nAnswer:"
    response = model.generate_content(prompt)
    return response.text

def query_groq(text, question):
    """Query using Groq API"""
    prompt = f"Context: {text}\n\nQuestion: {question}\n\nAnswer:"
    completion = groq_client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="mixtral-8x7b-32768",
    )
    return completion.choices[0].message.content

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        try:
            # Process file and create index
            text = process_file(filepath)
            documents = [Document(text=text)]  # Corrected: Use named argument `text`
            
            # Create index directly with embedding model
            index = VectorStoreIndex.from_documents(
                documents,
                embed_model=embed_model
            )
            
            # Save index temporarily
            temp_index_path = os.path.join(tempfile.gettempdir(), f"index_{filename}")
            index.storage_context.persist(persist_dir=temp_index_path)
            
            return jsonify({
                'message': 'File uploaded successfully',
                'filename': filename,
                'index_path': temp_index_path
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500


@app.route('/query', methods=['POST'])
def query_document():
    data = request.json
    question = data.get('question')
    model_choice = data.get('model', 'gemini')  # Default to Gemini
    filename = data.get('filename')
    index_path = data.get('index_path')
    
    if not all([question, filename, index_path]):
        return jsonify({'error': 'Missing required parameters'}), 400
    
    try:
        # Load the document text
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        text = process_file(filepath)
        
        # Query the selected model
        if model_choice == 'gemini':
            response = query_gemini(text, question)
        else:
            response = query_groq(text, question)
        
        return jsonify({'response': response})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
